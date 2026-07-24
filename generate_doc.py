import sys
import os
import zlib
import base64
import urllib.request
import time

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_kroki_url(text, diagram_type='mermaid'):
    compressed = zlib.compress(text.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('ascii')
    return f"https://kroki.io/{diagram_type}/png/{encoded}"

def download_image(url, filename, retries=3):
    for attempt in range(retries):
        try:
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'})
            with urllib.request.urlopen(req, timeout=30) as response, open(filename, 'wb') as out_file:
                out_file.write(response.read())
            return True
        except Exception as e:
            print(f"  Attempt {attempt+1} failed for {filename}: {e}")
            time.sleep(2)
    print(f"  SKIP: Could not download {filename}")
    return False

def add_diagram(doc, title, mermaid_code, filename, width=6.0):
    url = get_kroki_url(mermaid_code.strip(), 'mermaid')
    print(f"Generating: {title} -> {filename}")
    if download_image(url, filename):
        doc.add_picture(filename, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f'Hình: {title}')
        run.italic = True
        run.font.size = Pt(10)
        return True
    return False

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    return table

def add_steps(doc, steps):
    """Add numbered steps as a formatted list"""
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f'Bước {i}: ')
        run_num.bold = True
        p.add_run(step)

def main():
    doc = docx.Document()
    
    # ==========================================
    # CHƯƠNG 2
    # ==========================================
    doc.add_heading('CHƯƠNG 2 - ĐẶC TẢ YÊU CẦU', level=1)
    
    # --- 2.1 ---
    doc.add_heading('2.1. Khảo sát hiện trạng và bài toán đặt ra', level=2)
    doc.add_paragraph(
        'Trong bối cảnh thương mại điện tử phát triển mạnh mẽ tại Việt Nam, nhu cầu mua sắm trực tuyến ngày càng tăng cao. '
        'Tuy nhiên, nhiều cửa hàng nhỏ và vừa vẫn chưa có một hệ thống quản lý bán hàng toàn diện, '
        'dẫn đến việc quản lý thủ công gây ra nhiều sai sót và tốn kém thời gian.'
    )
    doc.add_paragraph(
        'Bài toán đặt ra là xây dựng một website thương mại điện tử hoàn chỉnh với hai phân hệ chính:'
    )
    doc.add_paragraph('• Phân hệ Khách hàng (Client): Cho phép người dùng đăng ký, đăng nhập, duyệt sản phẩm theo danh mục, '
        'tìm kiếm và lọc sản phẩm, thêm vào giỏ hàng, đặt hàng, thanh toán trực tuyến (MoMo, chuyển khoản ngân hàng qua VietQR), '
        'theo dõi đơn hàng, quản lý địa chỉ giao hàng, lưu sản phẩm yêu thích, xuất hoá đơn PDF, '
        'và tương tác với trợ lý AI tư vấn mua sắm.')
    doc.add_paragraph('• Phân hệ Quản trị (Admin): Cho phép quản trị viên quản lý toàn bộ danh mục, sản phẩm (bao gồm gallery ảnh), '
        'đơn hàng (cập nhật trạng thái, xác nhận thanh toán), người dùng (phân quyền Admin/User), banner quảng cáo, '
        'section trang chủ, cài đặt hệ thống (logo, thông tin ngân hàng), và xem báo cáo thống kê doanh thu.')
    
    # --- 2.2 ---
    doc.add_heading('2.2. Phân tích yêu cầu chức năng', level=2)
    
    doc.add_heading('2.2.1. Yêu cầu chức năng phía Khách hàng', level=3)
    add_table(doc, ['Chức năng', 'Mô tả chi tiết'], [
        ('Đăng ký / Đăng nhập', 'Người dùng đăng ký tài khoản với thông tin cá nhân (họ tên, email, username, số điện thoại, địa chỉ). Đăng nhập bằng email hoặc username. Hệ thống tự động tạo địa chỉ mặc định khi đăng ký.'),
        ('Duyệt sản phẩm', 'Xem danh sách sản phẩm với phân trang (12 sản phẩm/trang), lọc theo danh mục, khoảng giá (min/max), từ khóa tìm kiếm, và sắp xếp (giá tăng/giảm, tên A-Z/Z-A, mới nhất).'),
        ('Chi tiết sản phẩm', 'Xem thông tin chi tiết sản phẩm, gallery ảnh phụ với zoom modal, mô tả sản phẩm có thể thu gọn/mở rộng, sản phẩm liên quan cùng danh mục.'),
        ('Giỏ hàng', 'Thêm sản phẩm vào giỏ hàng (lưu Session), cập nhật số lượng, xóa sản phẩm, xem tổng tiền. Hỗ trợ nút "Mua ngay" chuyển thẳng đến trang thanh toán.'),
        ('Đặt hàng & Thanh toán', 'Chọn địa chỉ giao hàng, chọn phương thức thanh toán (COD, MoMo, Chuyển khoản). Hệ thống tự sinh mã đơn hàng dạng OD + timestamp, gửi email xác nhận đơn hàng tự động.'),
        ('Theo dõi đơn hàng', 'Xem lịch sử đơn hàng với phân trang, lọc theo trạng thái (chờ xử lý, đang xử lý, đang giao, hoàn thành, đã hủy), xem chi tiết từng đơn, xuất hoá đơn PDF.'),
        ('Quản lý địa chỉ', 'Thêm/sửa/xóa nhiều địa chỉ giao hàng, đặt địa chỉ mặc định. Tích hợp API GHN để lấy danh sách Tỉnh/Thành → Quận/Huyện → Phường/Xã tự động.'),
        ('Sản phẩm yêu thích', 'Lưu/bỏ lưu sản phẩm yêu thích vào LocalStorage (không cần đăng nhập), xem trang danh sách sản phẩm yêu thích riêng. Icon trái tim hiển thị trên mọi card sản phẩm.'),
        ('Trợ lý AI tư vấn', 'Chatbot AI tư vấn mua sắm tích hợp Google Gemini API, gợi ý sản phẩm dựa trên ngữ cảnh hội thoại và dữ liệu thực từ CSDL. Lịch sử chat lưu LocalStorage, giới hạn 20 lượt/phiên.'),
        ('Đổi mật khẩu & Hồ sơ', 'Xem và cập nhật thông tin cá nhân (họ tên, SĐT), đổi mật khẩu với xác minh mật khẩu cũ.'),
    ])
    
    doc.add_heading('2.2.2. Yêu cầu chức năng phía Quản trị viên', level=3)
    add_table(doc, ['Chức năng', 'Mô tả chi tiết'], [
        ('Dashboard thống kê', 'Xem tổng số đơn, doanh thu, số khách hàng theo khoảng thời gian tùy chọn. Biểu đồ cột doanh thu, biểu đồ tròn trạng thái đơn hàng, biểu đồ đường số đơn/ngày, top 5 sản phẩm bán chạy. Xuất báo cáo PDF.'),
        ('Quản lý Danh mục', 'CRUD danh mục sản phẩm, hỗ trợ cấu trúc danh mục cha-con (parent_id tự tham chiếu), upload ảnh danh mục, tự động tạo slug từ tên.'),
        ('Quản lý Sản phẩm', 'CRUD sản phẩm, upload ảnh chính + gallery nhiều ảnh phụ (bảng product_images), bật/tắt trạng thái hiển thị, đánh dấu sản phẩm nổi bật (is_featured), phân trang 10 SP/trang.'),
        ('Quản lý Đơn hàng', 'Xem danh sách đơn hàng với bộ lọc (ngày bắt đầu/kết thúc, trạng thái, ID khách hàng). Cập nhật trạng thái đơn: pending → processing → shipping → completed/cancelled. Xác nhận/hủy thanh toán chuyển khoản (markPaid/markUnpaid).'),
        ('Quản lý Người dùng', 'Xem danh sách người dùng (email/SĐT được che bảo mật bằng dấu *, bấm icon mắt để hiện). Nâng/hạ quyền Admin (bảo vệ admin cuối cùng), xóa tài khoản.'),
        ('Quản lý Banner', 'CRUD banner quảng cáo hiển thị trên trang chủ, upload ảnh hoặc nhập URL, sắp xếp thứ tự hiển thị, bật/tắt hiển thị.'),
        ('Quản lý Section trang chủ', 'Tùy chỉnh các khu vực hiển thị sản phẩm trên trang chủ: Sản phẩm nổi bật (type=1), Sản phẩm mới (type=2), Theo danh mục (type=3). Kéo thả sắp xếp thứ tự, bật/tắt.'),
        ('Cài đặt hệ thống', 'Cấu hình tên website, logo, mô tả. Cấu hình thanh toán: bank_bin, bank_name, bank_account_no, bank_account_name cho VietQR.'),
    ])
    
    # --- 2.3 ---
    doc.add_heading('2.3. Yêu cầu phi chức năng', level=2)
    doc.add_paragraph('• Hiệu năng: Tốc độ tải trang nhanh, sử dụng Eager Loading để tối ưu truy vấn N+1, phân trang dữ liệu toàn bộ danh sách.')
    doc.add_paragraph('• Bảo mật: Mã hóa mật khẩu bcrypt, bảo vệ CSRF token trên mọi form POST, che giấu thông tin nhạy cảm (email/SĐT) trong trang quản trị, middleware phân quyền Admin, xác thực chữ ký HMAC SHA256 từ MoMo.')
    doc.add_paragraph('• Đa ngôn ngữ: Hỗ trợ 2 ngôn ngữ Tiếng Việt và English thông qua Laravel Localization (resources/lang/vi, en).')
    doc.add_paragraph('• Giao diện: Responsive trên mobile/tablet/desktop, thiết kế hiện đại sử dụng Bootstrap 5, hiệu ứng animation, particles background, chatbot dạng floating panel.')
    doc.add_paragraph('• Khả năng mở rộng: Kiến trúc MVC rõ ràng, tách biệt Controller Admin/Client, Service layer cho tích hợp bên ngoài (MomoService, GHNApiService).')
    
    # =============================================
    # 2.4 Use Case Khách hàng (Diagram 1)
    # =============================================
    doc.add_heading('2.4. Mô hình Use Case (Tổng quát)', level=2)
    doc.add_heading('2.4.1. Use Case - Phía Khách hàng', level=3)
    doc.add_paragraph('Sơ đồ Use Case dưới đây thể hiện toàn bộ các chức năng mà khách hàng có thể thực hiện trên hệ thống:')
    
    uc_client = """
flowchart LR
    KH((Khach hang))
    subgraph HT [Phan he Khach hang]
        UC1(Dang ky tai khoan)
        UC2(Dang nhap / Dang xuat)
        UC3(Duyet san pham)
        UC4(Tim kiem va loc SP)
        UC5(Xem chi tiet SP)
        UC6(Them vao gio hang)
        UC7(Dat hang va Thanh toan)
        UC8(Theo doi don hang)
        UC9(Quan ly dia chi)
        UC10(Doi mat khau)
        UC11(Luu SP yeu thich)
        UC12(Chat voi AI tu van)
        UC13(Xuat hoa don PDF)
    end
    KH --- UC1
    KH --- UC2
    KH --- UC3
    KH --- UC4
    KH --- UC5
    KH --- UC6
    KH --- UC7
    KH --- UC8
    KH --- UC9
    KH --- UC10
    KH --- UC11
    KH --- UC12
    KH --- UC13
"""
    add_diagram(doc, 'Use Case - Phía Khách hàng', uc_client, 'd01.png')
    
    doc.add_paragraph('Mô tả chi tiết từng Use Case:')
    add_table(doc, ['Use Case', 'Actor', 'Mô tả', 'Tiền điều kiện', 'Luồng chính'], [
        ('Đăng ký tài khoản', 'Khách hàng', 'Tạo tài khoản mới với họ tên, email, username, SĐT, mật khẩu, địa chỉ.', 'Chưa có tài khoản', 'Nhập thông tin → Validate → Tạo User + Address mặc định → Tự động đăng nhập → Redirect trang chủ.'),
        ('Đăng nhập / Đăng xuất', 'Khách hàng', 'Xác thực danh tính bằng email hoặc username + mật khẩu.', 'Đã có tài khoản', 'Nhập email/username + password → Auth::attempt() → Regenerate session → Redirect trang chủ.'),
        ('Duyệt sản phẩm', 'Khách hàng', 'Xem danh sách sản phẩm đang hoạt động, phân trang 12 SP/trang.', 'Không cần đăng nhập', 'Truy cập /products → Controller query Product active → Paginate → Render danh sách.'),
        ('Tìm kiếm và lọc SP', 'Khách hàng', 'Lọc SP theo keyword (tên/mô tả), danh mục, giá min/max, sắp xếp.', 'Không cần đăng nhập', 'Nhập keyword/chọn filter → GET /products?q=...&category=...&min_price=... → Query với WHERE/ORDER BY → Render kết quả.'),
        ('Xem chi tiết SP', 'Khách hàng', 'Xem thông tin đầy đủ: ảnh chính + gallery, giá, tồn kho, mô tả, SP liên quan.', 'Không cần đăng nhập', 'Click vào SP → GET /products/{slug} → Load Product + images + 4 related products → Render.'),
        ('Thêm vào giỏ hàng', 'Khách hàng', 'Thêm SP vào giỏ hàng Session với số lượng tùy chọn.', 'Đã đăng nhập', 'Click "Thêm vào giỏ" → POST /cart/add → Cập nhật session cart → Trả về JSON + Toast thông báo + Cập nhật badge.'),
        ('Đặt hàng & Thanh toán', 'Khách hàng', 'Tạo đơn hàng từ giỏ hàng, chọn địa chỉ và phương thức thanh toán.', 'Giỏ hàng không rỗng', 'Chọn địa chỉ + payment method → POST /orders → Tạo Order + OrderItems → Xóa cart → Gửi email → Redirect trang thanh toán.'),
        ('Theo dõi đơn hàng', 'Khách hàng', 'Xem lịch sử đơn hàng, lọc theo trạng thái, xem chi tiết.', 'Đã đăng nhập', 'GET /orders → Query orders theo user_id + filter status → Paginate → Render danh sách. Click đơn → GET /orders/{id} → Chi tiết.'),
        ('Quản lý địa chỉ', 'Khách hàng', 'CRUD địa chỉ giao hàng, đặt mặc định. Chọn Tỉnh/Huyện/Xã qua API GHN.', 'Đã đăng nhập', 'GET /addresses → Danh sách. POST /addresses → Tạo mới (gọi GHN API lấy tỉnh/huyện/xã). PUT → Sửa. DELETE → Xóa (tự chuyển default).'),
        ('Đổi mật khẩu', 'Khách hàng', 'Thay đổi mật khẩu với xác minh mật khẩu hiện tại.', 'Đã đăng nhập', 'Nhập password cũ + password mới → Hash::check() xác minh → bcrypt hash password mới → Update DB.'),
        ('Lưu SP yêu thích', 'Khách hàng', 'Toggle trái tim trên card SP, lưu vào LocalStorage trình duyệt.', 'Không cần đăng nhập', 'Click icon trái tim → toggleWishlist() JS → Lưu/xóa khỏi localStorage → Cập nhật badge + icon trái tim.'),
        ('Chat với AI tư vấn', 'Khách hàng', 'Hỏi đáp với chatbot AI về sản phẩm, nhận gợi ý SP phù hợp.', 'Không cần đăng nhập', 'Nhập câu hỏi → POST /ai/consult → extractKeywords → Query DB tìm SP liên quan → Gọi Gemini API → Trả về reply + suggested products.'),
        ('Xuất hoá đơn PDF', 'Khách hàng', 'Tải xuống file PDF hoá đơn cho đơn hàng đã đặt.', 'Đã đăng nhập, có đơn hàng', 'Click "Tải hoá đơn" → GET /orders/{id}/invoice → Render Blade template → dompdf stream download.'),
    ])
    
    # =============================================
    # 2.4.2 Use Case Admin (Diagram 2)
    # =============================================
    doc.add_heading('2.4.2. Use Case - Phía Quản trị viên', level=3)
    doc.add_paragraph('Sơ đồ Use Case dưới đây thể hiện toàn bộ các chức năng của quản trị viên:')
    
    uc_admin = """
flowchart LR
    AD((Quan tri vien))
    subgraph HT [Phan he Quan tri]
        A1(Xem Dashboard thong ke)
        A2(Quan ly Danh muc)
        A3(Quan ly San pham)
        A4(Quan ly Don hang)
        A5(Cap nhat trang thai DH)
        A6(Xac nhan thanh toan)
        A7(Quan ly Nguoi dung)
        A8(Phan quyen Admin)
        A9(Quan ly Banner)
        A10(Quan ly Section trang chu)
        A11(Cai dat he thong)
        A12(Xuat bao cao PDF)
    end
    AD --- A1
    AD --- A2
    AD --- A3
    AD --- A4
    AD --- A5
    AD --- A6
    AD --- A7
    AD --- A8
    AD --- A9
    AD --- A10
    AD --- A11
    AD --- A12
"""
    add_diagram(doc, 'Use Case - Phía Quản trị viên', uc_admin, 'd02.png')
    
    doc.add_paragraph('Mô tả chi tiết từng Use Case phía Quản trị viên:')
    add_table(doc, ['Use Case', 'Mô tả', 'Luồng chính'], [
        ('Xem Dashboard thống kê', 'Trang tổng quan hiển thị các chỉ số kinh doanh theo khoảng thời gian.', 'Chọn ngày bắt đầu/kết thúc → Lọc → Query tổng đơn (count), doanh thu (sum total_amount), số user (count), top 5 SP bán chạy (JOIN order_items GROUP BY product_id) → Render biểu đồ Chart.js.'),
        ('Quản lý Danh mục', 'CRUD danh mục, hỗ trợ cha-con, upload ảnh, tự động slug.', 'Index: phân trang 10/trang, filter keyword + parent_id. Create: nhập name → Str::slug() tạo slug → upload ảnh → save. Edit/Update: sửa thông tin + thay ảnh. Delete: xóa ảnh trên disk + xóa record.'),
        ('Quản lý Sản phẩm', 'CRUD sản phẩm, gallery ảnh, toggle trạng thái.', 'Index: phân trang 10/trang, eager load category. Create: nhập thông tin + upload ảnh chính + upload nhiều ảnh gallery (ProductImage). Toggle: đảo is_active. DeleteImage: xóa ảnh gallery đơn lẻ.'),
        ('Quản lý Đơn hàng', 'Xem danh sách, filter, chi tiết, cập nhật trạng thái.', 'Index: filter date_from/date_to + status + user_id, eager load user + shippingAddress, paginate 10. Show: eager load items.product. UpdateStatus: validate status enum → update.'),
        ('Xác nhận thanh toán', 'Xác nhận/hủy thanh toán chuyển khoản ngân hàng.', 'markPaid: ghi timestamp + admin_id + transaction_id vào payment_payload → payment_status=paid → nếu status=pending thì chuyển processing. markUnpaid: payment_status=failed, xóa paid_at.'),
        ('Quản lý Người dùng', 'Xem danh sách, che bảo mật thông tin, phân quyền.', 'Index: paginate 10, đếm admin count. Show: chi tiết. ToggleAdmin: đảo is_admin (bảo vệ admin cuối cùng không cho hạ). Destroy: xóa user (không cho xóa admin cuối).'),
        ('Quản lý Banner', 'CRUD banner quảng cáo trên trang chủ.', 'Index: orderBy order. Create/Store: validate + upload ảnh hoặc nhập URL. Edit/Update: thay ảnh nếu có. Destroy: xóa record.'),
        ('Quản lý Section trang chủ', 'Tùy chỉnh khu vực hiển thị SP trên trang chủ.', 'Index: orderBy order. Create: chọn type (1=nổi bật, 2=mới, 3=theo danh mục) + chọn categories. UpdateOrder: AJAX cập nhật thứ tự hàng loạt. ToggleActive: bật/tắt hiển thị.'),
        ('Cài đặt hệ thống', 'Cấu hình website và thông tin ngân hàng.', 'Index: load all settings keyed by key. Update: validate + updateOrCreate cho site_name, site_description, logo (upload file), bank_bin, bank_name, bank_account_no, bank_account_name.'),
        ('Xuất báo cáo PDF', 'Xuất báo cáo doanh thu ra file PDF.', 'GET /admin/dashboard/export-pdf → tính lại metrics + top 10 SP → render Blade template admin.dashboard.pdf → Barryvdh DomPDF stream download.'),
    ])
    
    # =============================================
    # 2.5 FDD (Diagram 3)
    # =============================================
    doc.add_heading('2.5. Sơ đồ phân rã chức năng (Functional Decomposition Diagram)', level=2)
    doc.add_paragraph(
        'Sơ đồ phân rã chức năng thể hiện cách hệ thống được chia nhỏ thành các phân hệ và chức năng con. '
        'Hệ thống gồm 3 phân hệ chính: Phân hệ Khách hàng, Phân hệ Quản trị, và Phân hệ Thanh toán.'
    )
    
    fdd = """
graph TD
    A[He thong Thuong mai dien tu] --> B[Phan he Khach hang]
    A --> C[Phan he Quan tri]
    A --> D[Phan he Thanh toan]
    B --> B1[Quan ly tai khoan]
    B --> B2[Duyet va Tim kiem SP]
    B --> B3[Gio hang]
    B --> B4[Dat hang]
    B --> B5[Theo doi don hang]
    B --> B6[Quan ly dia chi]
    B --> B7[Wishlist - Yeu thich]
    B --> B8[AI tu van]
    C --> C1[QL San pham va Danh muc]
    C --> C2[QL Don hang]
    C --> C3[QL Nguoi dung]
    C --> C4[QL Banner va Section]
    C --> C5[Thong ke va Bao cao]
    C --> C6[Cai dat he thong]
    D --> D1[COD - Thanh toan khi nhan]
    D --> D2[MoMo e-Wallet]
    D --> D3[Chuyen khoan VietQR]
"""
    add_diagram(doc, 'Sơ đồ phân rã chức năng', fdd, 'd03.png')
    
    doc.add_paragraph('Mô tả chi tiết từng nhánh chức năng:')
    doc.add_heading('Phân hệ Khách hàng', level=4)
    add_table(doc, ['Chức năng', 'Controller xử lý', 'Mô tả nghiệp vụ'], [
        ('Quản lý tài khoản', 'RegisterController, LoginController, Client\\UserController', 'Đăng ký (tạo User + Address mặc định), đăng nhập (email/username), đăng xuất, xem/sửa hồ sơ, đổi mật khẩu.'),
        ('Duyệt & Tìm kiếm SP', 'Client\\ProductController', 'index(): phân trang 12 SP, filter theo q (LIKE name/description), category slug, min_price, max_price, sort. show(): chi tiết SP + 4 related products.'),
        ('Giỏ hàng', 'Client\\CartController', 'add(): thêm/cập nhật session cart. update(): đổi quantity. remove(): xóa item. checkout(): load cart + 3 địa chỉ user → render form thanh toán.'),
        ('Đặt hàng', 'Client\\OrderController', 'store(): validate address_id → tính total → tạo Order (order_number = OD + time) → tạo OrderItems → xóa session cart → gửi Mail OrderPlaced → redirect theo payment method.'),
        ('Theo dõi đơn hàng', 'Client\\OrderController', 'index(): paginate 10, filter status. show(): authorize user_id, eager load items.product. exportInvoice(): render Blade → dompdf stream PDF.'),
        ('Quản lý địa chỉ', 'AddressController + GHNApiService', 'CRUD + setDefault. getProvinces/getDistricts/getWards: proxy GHN API. Khi xóa địa chỉ default → tự động chuyển default cho địa chỉ mới nhất.'),
        ('Wishlist', 'JavaScript (wishlist.js)', 'toggleWishlist(id, name, price, image, url): toggle LocalStorage. updateBadge(): cập nhật số lượng. renderWishlistPage(): render danh sách.'),
        ('AI tư vấn', 'Client\\AiAssistantController', 'consult(): extractKeywords → findRelatedProducts (LIKE) → getFeaturedProducts → buildSystemInstruction (context) → callGeminiWithFallback → return reply + suggested products.'),
    ])
    
    doc.add_heading('Phân hệ Quản trị', level=4)
    add_table(doc, ['Chức năng', 'Controller xử lý', 'Mô tả nghiệp vụ'], [
        ('QL Sản phẩm & Danh mục', 'Admin\\ProductController, Admin\\CategoryController', 'CRUD đầy đủ. Product: upload ảnh chính + gallery (ProductImage). Category: self-reference parent_id, auto slug. Toggle is_active.'),
        ('QL Đơn hàng', 'Admin\\OrderController', 'index: filter date/status/user_id, eager load. updateStatus: validate enum. markPaid: ghi admin_id + timestamp vào payment_payload JSON. markUnpaid: revert.'),
        ('QL Người dùng', 'Admin\\UserController', 'index: paginate + admin count. toggleAdmin: đảo is_admin (bảo vệ admin cuối). destroy: xóa user (bảo vệ admin cuối).'),
        ('QL Banner & Section', 'Admin\\BannerController, Admin\\HomeSectionController', 'Banner: CRUD + upload. Section: CRUD + updateOrder (AJAX bulk) + toggleActive. Type: 1=featured, 2=new, 3=by category.'),
        ('Thống kê & Báo cáo', 'Admin\\DashboardController', 'index: count orders, sum revenue, count users, top 5 SP (JOIN order_items). exportPdf: top 10 SP → dompdf download.'),
        ('Cài đặt hệ thống', 'Admin\\SettingController', 'index: load all key-value. update: updateOrCreate cho site_name, site_description, logo, bank_bin, bank_name, bank_account_no, bank_account_name.'),
    ])
    
    doc.add_heading('Phân hệ Thanh toán', level=4)
    add_table(doc, ['Phương thức', 'Controller xử lý', 'Mô tả quy trình'], [
        ('COD', 'Client\\OrderController', 'Không xử lý thanh toán online. Sau khi tạo đơn → redirect thẳng trang chi tiết đơn hàng. Admin xác nhận thanh toán khi giao hàng thành công.'),
        ('MoMo e-Wallet', 'Client\\PaymentController + MomoService', 'momoShow: gọi MoMo API tạo payment → nhận payUrl + qrCodeUrl → lưu payment_payload. momoIpn: webhook server-to-server, verify HMAC signature → update payment_status. momoReturn: redirect user về.'),
        ('Chuyển khoản VietQR', 'Client\\PaymentController + Setting Model', 'bankShow: lấy thông tin NH từ settings → tạo URL QR VietQR với amount + order_number → render QR. bankNotify: khách xác nhận đã chuyển. Admin markPaid thủ công.'),
    ])
    
    # ==========================================
    # CHƯƠNG 3
    # ==========================================
    doc.add_heading('CHƯƠNG 3 - THIẾT KẾ VÀ CÀI ĐẶT GIẢI PHÁP', level=1)
    
    # =============================================
    # 3.1 Kiến trúc (Diagram 4)
    # =============================================
    doc.add_heading('3.1. Thiết kế kiến trúc tổng thể (System Architecture)', level=2)
    doc.add_paragraph(
        'Hệ thống được xây dựng theo mô hình MVC (Model-View-Controller) trên nền tảng Laravel Framework, '
        'gồm 4 tầng chính: Client Tier, Application Tier, Data Tier, và External Services.'
    )
    
    arch = """
graph TD
    subgraph Client [Client Tier]
        BR[Web Browser]
        JS[JavaScript - jQuery - AJAX]
        LS[LocalStorage]
    end
    subgraph App [Application Tier - Laravel MVC]
        RT[Routes - web.php]
        MW[Middleware - Auth - Admin]
        CT[Controllers - Admin - Client - Auth]
        MD[Eloquent Models - 10 Models]
        BL[Blade Views - layouts - admin - client]
        SV[Services - MomoService - GHNApiService]
    end
    subgraph Data [Data Tier]
        DB[(MySQL 8.0 - 10 tables)]
        FS[File Storage - public/storage]
    end
    subgraph Ext [External APIs]
        MOMO[MoMo Payment v2]
        GHN[GHN Address API]
        GEMINI[Google Gemini AI]
        VIETQR[VietQR NAPAS247]
        SMTP[SMTP Mail]
    end
    BR --> RT
    RT --> MW
    MW --> CT
    CT --> MD
    CT --> BL
    BL --> BR
    MD --> DB
    CT --> FS
    CT --> SV
    SV --> MOMO
    SV --> GHN
    CT --> GEMINI
    CT --> VIETQR
    CT --> SMTP
    JS --> BR
    LS --> BR
"""
    add_diagram(doc, 'Kiến trúc tổng thể hệ thống', arch, 'd04.png')
    
    doc.add_paragraph('Mô tả chi tiết luồng xử lý qua từng tầng:')
    add_steps(doc, [
        'Client Tier: Người dùng tương tác qua Web Browser. JavaScript (jQuery) xử lý AJAX request (thêm giỏ hàng, AI chat). LocalStorage lưu wishlist và lịch sử chat AI.',
        'Routes (web.php): Định tuyến URL đến Controller tương ứng. Chia thành 4 nhóm: Public (không cần đăng nhập), Guest (login/register), Auth (yêu cầu đăng nhập), Admin (yêu cầu đăng nhập + quyền admin).',
        'Middleware: Auth middleware kiểm tra session đăng nhập. Admin middleware kiểm tra cột is_admin=true trong bảng users. CSRF middleware bảo vệ mọi form POST.',
        'Controllers: 17 controller chia thành 4 namespace: Admin (8 controller), Client (6 controller), Auth (2 controller), Root (2 controller). Mỗi controller chứa logic nghiệp vụ và gọi Model.',
        'Eloquent Models: 10 Model (User, Category, Product, ProductImage, Order, OrderItem, Address, Banner, HomeSection, Setting) ánh xạ đến 10 bảng MySQL. Định nghĩa relationships (belongsTo, hasMany), casts, accessors.',
        'Blade Views: Template engine render HTML. Chia thành layouts/ (eshopper.blade.php cho client, admin.blade.php cho admin), admin/ (8 module), client/ (8 module).',
        'Services: MomoService xử lý tạo payment request + verify signature. GHNApiService proxy API lấy danh sách tỉnh/huyện/xã.',
        'Data Tier: MySQL 8.0 lưu trữ 10 bảng chính. File Storage (public/storage) lưu ảnh sản phẩm, danh mục, banner, logo.',
        'External APIs: MoMo Payment v2 (thanh toán ví điện tử), GHN API (địa chỉ), Google Gemini (AI chatbot), VietQR (QR chuyển khoản), SMTP (gửi email).',
    ])
    
    # =============================================
    # 3.2 Sequence Diagrams
    # =============================================
    doc.add_heading('3.2. Thiết kế Sơ đồ tuần tự (Sequence Diagram)', level=2)
    
    # --- Diagram 5: Đăng ký ---
    doc.add_heading('3.2.1. Sơ đồ tuần tự - Đăng ký tài khoản', level=3)
    doc.add_paragraph('Mô tả luồng xử lý khi khách hàng đăng ký tài khoản mới trên hệ thống (RegisterController):')
    
    seq_register = """
sequenceDiagram
    actor User
    participant Browser
    participant RegisterController
    participant UserModel as User Model
    participant AddressModel as Address Model
    participant DB as MySQL

    User->>Browser: Truy cap /register
    Browser->>RegisterController: GET /register
    RegisterController-->>Browser: Hien thi form dang ky

    User->>Browser: Dien thong tin va Submit
    Browser->>RegisterController: POST /register
    RegisterController->>RegisterController: Validate du lieu
    RegisterController->>UserModel: Create User
    UserModel->>DB: INSERT INTO users
    DB-->>UserModel: Return user record
    RegisterController->>AddressModel: Create Address mac dinh
    AddressModel->>DB: INSERT INTO addresses
    RegisterController->>RegisterController: Auth login user
    RegisterController-->>Browser: Redirect to /home
    Browser-->>User: Trang chu da dang nhap
"""
    add_diagram(doc, 'Sơ đồ tuần tự - Đăng ký tài khoản', seq_register, 'd05.png')
    
    doc.add_paragraph('Mô tả chi tiết từng bước:')
    add_steps(doc, [
        'Người dùng truy cập URL /register trên trình duyệt.',
        'RegisterController trả về view auth.register chứa form đăng ký.',
        'Người dùng điền đầy đủ thông tin: first_name, last_name, email (unique), username (unique), phone, password, password_confirmation, address, province, district, ward.',
        'RegisterController validate dữ liệu: kiểm tra email/username chưa tồn tại, password tối thiểu 8 ký tự, password_confirmation khớp.',
        'Tạo bản ghi User trong bảng users: mật khẩu được mã hoá bằng bcrypt (Hash::make). Trường is_admin mặc định là false.',
        'Tạo bản ghi Address trong bảng addresses: gắn user_id vừa tạo, đặt is_default = true (địa chỉ mặc định đầu tiên).',
        'Gọi Auth::login($user) để tự động đăng nhập người dùng ngay sau khi đăng ký.',
        'Redirect về trang chủ /home với trạng thái đã đăng nhập.',
    ])
    
    # --- Diagram 6: Đăng nhập ---
    doc.add_heading('3.2.2. Sơ đồ tuần tự - Đăng nhập', level=3)
    doc.add_paragraph('Mô tả luồng xử lý khi người dùng đăng nhập hệ thống (LoginController). Hệ thống hỗ trợ đăng nhập bằng cả email lẫn username:')
    
    seq_login = """
sequenceDiagram
    actor User
    participant Browser
    participant LoginController
    participant Auth as Auth Facade
    participant DB as MySQL

    User->>Browser: Truy cap /login
    Browser->>LoginController: GET /login
    LoginController-->>Browser: Hien thi form dang nhap

    User->>Browser: Nhap email hoac username va password
    Browser->>LoginController: POST /login
    LoginController->>LoginController: Xac dinh input la email hay username
    LoginController->>Auth: Auth attempt voi credentials
    Auth->>DB: SELECT FROM users WHERE email/username
    DB-->>Auth: User record

    alt Dung mat khau
        Auth-->>LoginController: true
        LoginController->>LoginController: session regenerate
        LoginController-->>Browser: Redirect to intended URL
        Browser-->>User: Trang chu
    else Sai mat khau
        Auth-->>LoginController: false
        LoginController-->>Browser: Redirect back voi loi
        Browser-->>User: Hien thi thong bao loi
    end
"""
    add_diagram(doc, 'Sơ đồ tuần tự - Đăng nhập', seq_login, 'd06.png')
    
    doc.add_paragraph('Mô tả chi tiết từng bước:')
    add_steps(doc, [
        'Người dùng truy cập /login, LoginController trả về view auth.login.',
        'Người dùng nhập thông tin vào trường "login" (có thể là email hoặc username) và "password".',
        'LoginController kiểm tra input: nếu chứa ký tự @ thì xác định là email, ngược lại là username.',
        'Gọi Auth::attempt() với credentials tương ứng (["email" => input, "password" => password] hoặc ["username" => input, "password" => password]).',
        'Laravel tự động truy vấn bảng users, tìm user theo email/username, sau đó so khớp password với hash bcrypt trong DB.',
        'Nếu đúng: regenerate session ID (chống session fixation), redirect đến URL trước đó (intended) hoặc trang chủ.',
        'Nếu sai: redirect quay lại trang login kèm thông báo lỗi "Sai thông tin đăng nhập".',
    ])
    
    # --- Diagram 7: Giỏ hàng ---
    doc.add_heading('3.2.3. Sơ đồ tuần tự - Thêm sản phẩm vào giỏ hàng', level=3)
    doc.add_paragraph('Mô tả luồng xử lý khi khách hàng thêm sản phẩm vào giỏ hàng (CartController). Giỏ hàng được lưu trong Laravel Session:')
    
    seq_cart = """
sequenceDiagram
    actor User
    participant Browser
    participant CartController
    participant Product as Product Model
    participant Session as Laravel Session

    User->>Browser: Click Them vao gio hang
    Browser->>CartController: POST /cart/add product_id qty
    CartController->>Product: Find product by ID
    Product-->>CartController: Product data

    CartController->>Session: Lay session cart hien tai
    Session-->>CartController: Cart array

    alt SP da co trong gio
        CartController->>CartController: Tang so luong
    else SP chua co
        CartController->>CartController: Them moi vao cart
    end

    CartController->>Session: Luu lai session cart
    CartController-->>Browser: JSON cart_count va message
    Browser-->>User: Toast thong bao va cap nhat badge
"""
    add_diagram(doc, 'Sơ đồ tuần tự - Thêm vào giỏ hàng', seq_cart, 'd07.png')
    
    doc.add_paragraph('Mô tả chi tiết từng bước:')
    add_steps(doc, [
        'Người dùng click nút "Thêm vào giỏ" trên card sản phẩm hoặc trang chi tiết, trình duyệt gửi AJAX POST đến /cart/add với product_id và quantity.',
        'CartController nhận request, truy vấn Product Model để lấy thông tin sản phẩm (tên, giá, ảnh).',
        'Đọc mảng cart hiện tại từ session: $cart = session("cart", []).',
        'Kiểm tra product_id đã tồn tại trong mảng cart chưa: Nếu đã có → cộng thêm quantity. Nếu chưa có → thêm phần tử mới với key = product_id, value = {name, price, image, quantity}.',
        'Lưu lại mảng cart vào session: session(["cart" => $cart]).',
        'Trả về JSON response chứa: cart_count (tổng số lượng SP trong giỏ), message (thông báo thành công). Nếu request có flag buy_now=1 → redirect thẳng đến /checkout.',
        'JavaScript phía client nhận response → hiển thị SweetAlert2 Toast ở góc trên bên phải + cập nhật số trên badge giỏ hàng ở header.',
    ])
    
    # --- Diagram 8: Đặt hàng ---
    doc.add_heading('3.2.4. Sơ đồ tuần tự - Đặt hàng (Checkout)', level=3)
    doc.add_paragraph('Mô tả luồng xử lý khi khách hàng hoàn tất đặt hàng (Client\\OrderController@store). Đây là luồng nghiệp vụ quan trọng nhất:')
    
    seq_order = """
sequenceDiagram
    actor User
    participant Browser
    participant OrderCtrl as OrderController
    participant Order as Order Model
    participant Item as OrderItem Model
    participant Mail as Mail Service
    participant Session
    participant DB as MySQL

    User->>Browser: Chon dia chi va phuong thuc TT
    Browser->>OrderCtrl: POST /orders
    OrderCtrl->>Session: Lay cart tu session
    Session-->>OrderCtrl: Cart array
    OrderCtrl->>OrderCtrl: Tinh total_amount

    OrderCtrl->>Order: Create Order
    Order->>DB: INSERT INTO orders
    DB-->>Order: Return order

    loop Moi san pham trong gio
        OrderCtrl->>Item: Create OrderItem
        Item->>DB: INSERT INTO order_items
    end

    OrderCtrl->>Session: Xoa cart session
    OrderCtrl->>Mail: Gui email OrderPlaced

    alt momo
        OrderCtrl-->>Browser: Redirect /payment/momo/id
    else bank_transfer
        OrderCtrl-->>Browser: Redirect /payment/bank/id
    else cod
        OrderCtrl-->>Browser: Redirect /orders/id
    end
"""
    add_diagram(doc, 'Sơ đồ tuần tự - Đặt hàng (Checkout)', seq_order, 'd08.png')
    
    doc.add_paragraph('Mô tả chi tiết từng bước:')
    add_steps(doc, [
        'Khách hàng ở trang /checkout chọn địa chỉ giao hàng (shipping_address_id), phương thức thanh toán (payment_method: cod/momo/bank_transfer), và nhập ghi chú (notes).',
        'Browser gửi POST /orders. OrderController validate: address_id phải thuộc về user hiện tại, payment_method phải hợp lệ.',
        'Lấy mảng cart từ session, kiểm tra giỏ hàng không rỗng. Nếu rỗng → redirect về /cart kèm lỗi.',
        'Tính tổng tiền total_amount = Σ(price × quantity) cho tất cả sản phẩm trong giỏ.',
        'Tạo bản ghi Order: order_number = "OD" + date("YmdHis") + random digits (đảm bảo unique), status = "pending", payment_status = "pending".',
        'Vòng lặp tạo OrderItem cho mỗi sản phẩm trong giỏ: ghi nhận order_id, product_id, quantity, price (giá tại thời điểm mua), subtotal = price × quantity.',
        'Xóa session cart: session()->forget("cart"). Giỏ hàng trở về trống.',
        'Gửi email xác nhận đơn hàng tự động qua Mail::to($user)->send(new OrderPlaced($order)).',
        'Redirect theo phương thức thanh toán: MoMo → /payment/momo/{order_id}, Chuyển khoản → /payment/bank/{order_id}, COD → /orders/{order_id} (trang chi tiết đơn).',
    ])
    
    # --- Diagram 9: MoMo ---
    doc.add_heading('3.2.5. Sơ đồ tuần tự - Thanh toán MoMo', level=3)
    doc.add_paragraph('Mô tả luồng thanh toán qua ví điện tử MoMo (PaymentController + MomoService). Quy trình bao gồm tạo payment request, hiển thị QR, và xử lý IPN callback:')
    
    seq_momo = """
sequenceDiagram
    actor User
    participant Browser
    participant PayCtrl as PaymentController
    participant Svc as MomoService
    participant MoMo as MoMo Gateway
    participant Order as Order Model
    participant DB as MySQL

    Browser->>PayCtrl: GET /payment/momo/order_id
    PayCtrl->>Svc: createPayment order
    Svc->>MoMo: POST Create Payment
    MoMo-->>Svc: payUrl va qrCodeUrl
    Svc-->>PayCtrl: Payment data
    PayCtrl->>Order: Luu payment_payload
    PayCtrl-->>Browser: Hien thi QR code

    User->>MoMo: Quet QR thanh toan

    MoMo->>PayCtrl: IPN Callback server-to-server
    PayCtrl->>PayCtrl: Verify HMAC SHA256
    PayCtrl->>Order: payment_status = paid
    Order->>DB: UPDATE orders

    MoMo->>Browser: Redirect /payment/momo/return
    PayCtrl-->>Browser: Ket qua thanh toan
    Browser-->>User: Thanh toan thanh cong
"""
    add_diagram(doc, 'Sơ đồ tuần tự - Thanh toán MoMo', seq_momo, 'd09.png')
    
    doc.add_paragraph('Mô tả chi tiết từng bước:')
    add_steps(doc, [
        'Sau khi đặt hàng với payment_method = "momo", khách được redirect đến /payment/momo/{order_id}.',
        'PaymentController gọi MomoService::createPayment($order): tạo request body gồm partnerCode, amount, orderId, requestId, returnUrl, ipnUrl, và ký HMAC SHA256.',
        'MomoService gửi POST request đến MoMo API endpoint. MoMo trả về payUrl (link thanh toán), qrCodeUrl (ảnh QR code), và requestId.',
        'Lưu toàn bộ response payload vào cột payment_payload (JSON) của order. Render view hiển thị QR code và link thanh toán.',
        'Khách hàng quét QR bằng app MoMo hoặc click link payUrl để thanh toán. Browser polling /payment/momo/{order_id}/status mỗi vài giây để check trạng thái.',
        'MoMo gửi IPN callback (server-to-server POST) đến /payment/momo/ipn. PaymentController verify chữ ký HMAC SHA256 để đảm bảo request hợp lệ.',
        'Nếu resultCode = 0 (thành công): update payment_status = "paid", paid_at = now(), status chuyển sang "processing". Nếu thất bại: payment_status = "failed".',
        'MoMo redirect trình duyệt khách về returnUrl = /payment/momo/return. PaymentController hiển thị kết quả thanh toán.',
    ])
    
    # --- Diagram 10: Bank ---
    doc.add_heading('3.2.6. Sơ đồ tuần tự - Thanh toán chuyển khoản (VietQR)', level=3)
    doc.add_paragraph('Mô tả luồng thanh toán qua chuyển khoản ngân hàng với mã QR VietQR/NAPAS247. Quy trình yêu cầu Admin xác nhận thủ công:')
    
    seq_bank = """
sequenceDiagram
    actor User
    participant Browser
    participant PayCtrl as PaymentController
    participant Setting as Setting Model
    participant Order as Order Model
    actor Admin

    Browser->>PayCtrl: GET /payment/bank/order_id
    PayCtrl->>Setting: Lay bank_bin account_no account_name
    PayCtrl->>PayCtrl: Tao VietQR URL voi amount va order_number
    PayCtrl-->>Browser: Hien thi QR chuyen khoan

    User->>User: Mo app NH quet QR chuyen khoan
    User->>Browser: Click Da chuyen khoan
    Browser->>PayCtrl: POST /payment/bank/order_id/notify
    PayCtrl->>Order: Ghi thoi diem khach xac nhan

    Admin->>Admin: Kiem tra tai khoan ngan hang
    Admin->>Order: POST admin/orders/id/mark-paid
    Order->>Order: payment_status = paid va status = processing
"""
    add_diagram(doc, 'Sơ đồ tuần tự - Thanh toán chuyển khoản (VietQR)', seq_bank, 'd10.png')
    
    doc.add_paragraph('Mô tả chi tiết từng bước:')
    add_steps(doc, [
        'Sau khi đặt hàng với payment_method = "bank_transfer", khách được redirect đến /payment/bank/{order_id}.',
        'PaymentController đọc thông tin ngân hàng từ bảng settings: bank_bin (mã BIN ngân hàng), bank_account_no (số tài khoản), bank_account_name (tên chủ tài khoản).',
        'Tạo URL QR code VietQR theo chuẩn NAPAS247: https://img.vietqr.io/image/{bank_bin}-{account_no}-...?amount={total_amount}&addInfo={order_number}. QR chứa thông tin chuyển khoản tự động (số tiền + nội dung).',
        'Render view hiển thị ảnh QR, thông tin tài khoản nhận tiền, số tiền cần chuyển, nội dung chuyển khoản = mã đơn hàng.',
        'Khách hàng mở app ngân hàng, quét mã QR (số tiền và nội dung tự điền), thực hiện chuyển khoản.',
        'Khách hàng click nút "Tôi đã chuyển khoản" trên giao diện → POST /payment/bank/{order_id}/notify → hệ thống ghi nhận thời điểm xác nhận vào payment_payload.',
        'Admin vào trang quản lý đơn hàng, kiểm tra tài khoản ngân hàng đã nhận tiền, click nút "Xác nhận đã thanh toán" → POST admin/orders/{id}/mark-paid.',
        'Hệ thống cập nhật: payment_status = "paid", paid_at = now(), ghi admin_id xác nhận, nếu status đang "pending" thì chuyển sang "processing".',
    ])
    
    # --- Diagram 11: AI ---
    doc.add_heading('3.2.7. Sơ đồ tuần tự - Trợ lý AI tư vấn sản phẩm', level=3)
    doc.add_paragraph('Mô tả luồng xử lý chatbot AI tư vấn (AiAssistantController). Chatbot sử dụng Google Gemini API kết hợp dữ liệu sản phẩm thực từ CSDL:')
    
    seq_ai = """
sequenceDiagram
    actor User
    participant Browser
    participant AiCtrl as AiAssistantController
    participant Product as Product Model
    participant Category as Category Model
    participant Gemini as Google Gemini API
    participant LS as LocalStorage

    User->>Browser: Nhap cau hoi vao chatbox
    Browser->>AiCtrl: POST /ai/consult
    AiCtrl->>AiCtrl: extractKeywords tu prompt
    AiCtrl->>Product: Query SP lien quan
    Product-->>AiCtrl: Related products
    AiCtrl->>Category: getCategorySummary
    Category-->>AiCtrl: Category list

    AiCtrl->>AiCtrl: buildSystemInstruction
    AiCtrl->>Gemini: POST generateContent
    Gemini-->>AiCtrl: AI response

    AiCtrl-->>Browser: JSON reply va suggested_products
    Browser->>LS: Luu lich su chat
    Browser-->>User: Hien thi tra loi va goi y SP
"""
    add_diagram(doc, 'Sơ đồ tuần tự - Trợ lý AI tư vấn sản phẩm', seq_ai, 'd11.png')
    
    doc.add_paragraph('Mô tả chi tiết từng bước:')
    add_steps(doc, [
        'Khách hàng click nút robot ở góc phải dưới → mở floating panel chatbot. Nhập câu hỏi (VD: "Tôi muốn mua điện thoại giá rẻ") → Submit.',
        'Browser gửi POST /ai/consult với body: prompt (câu hỏi), history (lịch sử chat trước đó), page (URL trang hiện tại để cung cấp ngữ cảnh).',
        'AiAssistantController gọi extractKeywords(): tách các từ khóa quan trọng từ câu hỏi (loại bỏ stop words tiếng Việt).',
        'Truy vấn Product Model: findRelatedProducts() với WHERE name LIKE %keyword% để tìm sản phẩm liên quan. getFeaturedProducts() lấy thêm sản phẩm nổi bật.',
        'Truy vấn Category Model: getCategorySummary() lấy danh sách toàn bộ danh mục + số lượng SP để cung cấp context cho AI.',
        'buildSystemInstruction(): xây dựng system prompt chứa: tên cửa hàng (từ settings), danh sách danh mục, danh sách sản phẩm thực (tên + giá + link) → AI sẽ trả lời dựa trên dữ liệu thực.',
        'Gọi Google Gemini API (callGeminiWithFallback): gửi system instruction + lịch sử chat + câu hỏi mới. Hỗ trợ fallback nhiều model (gemini-2.0-flash → gemini-1.5-flash → ...).',
        'Nhận response từ Gemini, trích xuất reply text. Kết hợp với suggested_products (sản phẩm liên quan từ DB) → trả JSON về browser.',
        'JavaScript render reply (hỗ trợ markdown: bold, italic, list, link) + hiển thị card sản phẩm gợi ý có tên + giá + link. Lưu lịch sử chat vào LocalStorage (key: ai_assistant_history_v1, giới hạn 20 lượt).',
    ])
    
    # --- Diagram 12: Admin Order ---
    doc.add_heading('3.2.8. Sơ đồ tuần tự - Admin quản lý đơn hàng', level=3)
    doc.add_paragraph('Mô tả luồng xử lý khi Admin xem và cập nhật trạng thái đơn hàng (Admin\\OrderController):')
    
    seq_admin = """
sequenceDiagram
    actor Admin
    participant Browser
    participant OrderCtrl as Admin OrderController
    participant Order as Order Model
    participant DB as MySQL

    Admin->>Browser: Truy cap /admin/orders
    Browser->>OrderCtrl: GET /admin/orders
    OrderCtrl->>Order: Query filter date status user_id
    Order->>DB: SELECT with Eager Load
    DB-->>Order: Order list
    OrderCtrl-->>Browser: Danh sach don hang

    Admin->>Browser: Click xem chi tiet
    Browser->>OrderCtrl: GET /admin/orders/id
    OrderCtrl->>Order: Eager load items.product
    OrderCtrl-->>Browser: Chi tiet don hang

    Admin->>Browser: Chon trang thai moi
    Browser->>OrderCtrl: POST /admin/orders/id/update-status
    OrderCtrl->>Order: Update status
    Order->>DB: UPDATE orders
    OrderCtrl-->>Browser: Redirect thanh cong
"""
    add_diagram(doc, 'Sơ đồ tuần tự - Admin quản lý đơn hàng', seq_admin, 'd12.png')
    
    doc.add_paragraph('Mô tả chi tiết từng bước:')
    add_steps(doc, [
        'Admin truy cập /admin/orders. Có thể sử dụng bộ lọc: ngày bắt đầu (date_from), ngày kết thúc (date_to), trạng thái đơn (status), ID khách hàng (user_id).',
        'OrderController query bảng orders với điều kiện filter, eager load relationship user và shippingAddress để hiển thị tên khách + địa chỉ. Phân trang 10 đơn/trang.',
        'Admin click vào đơn cụ thể → GET /admin/orders/{id}. Controller eager load items.product để hiển thị danh sách sản phẩm trong đơn (tên SP, số lượng, đơn giá, thành tiền).',
        'Admin chọn trạng thái mới từ dropdown (pending/processing/shipping/completed/cancelled) và submit.',
        'Controller validate trạng thái mới phải nằm trong enum cho phép, cập nhật cột status trong bảng orders, redirect kèm thông báo thành công.',
    ])
    
    # =============================================
    # 3.3 State Diagrams
    # =============================================
    doc.add_heading('3.3. Sơ đồ trạng thái (State Diagram)', level=2)
    
    # --- Diagram 13: Order State ---
    doc.add_heading('3.3.1. Sơ đồ trạng thái - Đơn hàng', level=3)
    doc.add_paragraph('Dựa trên cột status trong bảng orders với kiểu ENUM gồm 5 giá trị: pending, processing, shipping, completed, cancelled. Sơ đồ dưới đây mô tả vòng đời của một đơn hàng:')
    
    state_order = """
stateDiagram-v2
    [*] --> pending : Khach dat hang
    pending --> processing : Admin xac nhan
    pending --> cancelled : Khach hoac Admin huy
    processing --> shipping : Ban giao van chuyen
    processing --> cancelled : Admin huy
    shipping --> completed : Giao thanh cong
    shipping --> cancelled : Khach tu choi nhan
    completed --> [*]
    cancelled --> [*]
"""
    add_diagram(doc, 'Sơ đồ trạng thái - Đơn hàng', state_order, 'd13.png')
    
    doc.add_paragraph('Mô tả chi tiết từng trạng thái và chuyển đổi:')
    add_table(doc, ['Trạng thái', 'Ý nghĩa', 'Chuyển đổi tiếp theo'], [
        ('pending (Chờ xử lý)', 'Đơn hàng vừa được tạo, chưa được admin xác nhận. Đây là trạng thái khởi tạo mặc định.', '→ processing (admin xác nhận) hoặc → cancelled (khách/admin huỷ).'),
        ('processing (Đang xử lý)', 'Admin đã xác nhận đơn hàng, đang chuẩn bị hàng để giao. Nếu thanh toán MoMo/chuyển khoản, trạng thái này được kích hoạt sau khi xác nhận thanh toán.', '→ shipping (bàn giao cho đơn vị vận chuyển) hoặc → cancelled (admin huỷ).'),
        ('shipping (Đang giao)', 'Đơn hàng đã được bàn giao cho đơn vị vận chuyển và đang trên đường giao đến khách.', '→ completed (giao thành công) hoặc → cancelled (khách từ chối nhận hàng).'),
        ('completed (Hoàn thành)', 'Đơn hàng đã được giao thành công đến tay khách hàng. Đây là trạng thái kết thúc thành công.', 'Trạng thái cuối cùng, không chuyển đổi tiếp.'),
        ('cancelled (Đã huỷ)', 'Đơn hàng bị huỷ bởi khách hàng hoặc quản trị viên ở bất kỳ giai đoạn nào trước khi hoàn thành.', 'Trạng thái cuối cùng, không chuyển đổi tiếp.'),
    ])
    
    # --- Diagram 14: Payment State ---
    doc.add_heading('3.3.2. Sơ đồ trạng thái - Thanh toán', level=3)
    doc.add_paragraph('Dựa trên cột payment_status trong bảng orders với kiểu ENUM gồm 4 giá trị: pending, paid, failed, refunded:')
    
    state_pay = """
stateDiagram-v2
    [*] --> pending : Tao don hang
    pending --> paid : MoMo IPN OK hoac Admin xac nhan CK
    pending --> failed : Thanh toan that bai hoac Admin tu choi
    paid --> refunded : Admin hoan tien
    failed --> pending : Khach thanh toan lai
    paid --> [*]
    failed --> [*]
    refunded --> [*]
"""
    add_diagram(doc, 'Sơ đồ trạng thái - Thanh toán', state_pay, 'd14.png')
    
    doc.add_paragraph('Mô tả chi tiết từng trạng thái thanh toán:')
    add_table(doc, ['Trạng thái', 'Ý nghĩa', 'Trigger chuyển đổi'], [
        ('pending (Chờ thanh toán)', 'Đơn hàng vừa tạo, chưa có thanh toán. Mặc định cho mọi đơn.', 'COD: giữ pending cho đến khi giao hàng. MoMo: chờ IPN. Bank: chờ admin xác nhận.'),
        ('paid (Đã thanh toán)', 'Thanh toán thành công. Cột paid_at được ghi nhận timestamp.', 'MoMo: IPN callback resultCode=0 → auto update. Bank: Admin click markPaid → ghi admin_id + timestamp vào payment_payload.'),
        ('failed (Thất bại)', 'Thanh toán không thành công.', 'MoMo: IPN callback resultCode≠0. Bank: Admin click markUnpaid (từ chối). Xóa paid_at.'),
        ('refunded (Hoàn tiền)', 'Admin thực hiện hoàn tiền cho khách (xử lý ngoài hệ thống, hệ thống chỉ ghi nhận trạng thái).', 'Admin cập nhật thủ công khi đã hoàn tiền qua kênh ngân hàng.'),
    ])
    
    # =============================================
    # 3.4 ERD (Diagram 15)
    # =============================================
    doc.add_heading('3.4. Sơ đồ Quan hệ Cơ sở dữ liệu (ERD & Data Dictionary)', level=2)
    doc.add_heading('3.4.1. Sơ đồ ERD', level=3)
    doc.add_paragraph('Sơ đồ thực thể - quan hệ (Entity Relationship Diagram) thể hiện cấu trúc 10 bảng chính trong CSDL MySQL và các mối quan hệ giữa chúng:')
    
    erd = """
erDiagram
    users ||--o{ orders : places
    users ||--o{ addresses : has
    users {
        bigint id PK
        string first_name
        string last_name
        string email UK
        string username UK
        string phone
        string password
        boolean is_admin
    }
    categories ||--o{ products : contains
    categories ||--o{ categories : parent
    categories {
        bigint id PK
        string name
        string slug UK
        string image
        bigint parent_id FK
        boolean is_active
    }
    products ||--o{ order_items : has
    products ||--o{ product_images : gallery
    products {
        bigint id PK
        bigint category_id FK
        string name
        string slug UK
        decimal price
        integer stock
        string image
        boolean is_active
        boolean is_featured
    }
    product_images {
        bigint id PK
        bigint product_id FK
        string image_path
    }
    orders ||--o{ order_items : contains
    orders }o--|| addresses : ships_to
    orders {
        bigint id PK
        bigint user_id FK
        bigint shipping_address_id FK
        string order_number UK
        enum status
        enum payment_status
        string payment_method
        decimal total_amount
        json payment_payload
        timestamp paid_at
    }
    order_items {
        bigint id PK
        bigint order_id FK
        bigint product_id FK
        integer quantity
        decimal price
        decimal subtotal
    }
    addresses {
        bigint id PK
        bigint user_id FK
        string full_name
        string phone
        string address_col
        string province_name
        string district_name
        string ward_name
        boolean is_default
    }
    banners {
        bigint id PK
        string title
        string image
        boolean is_active
    }
    home_sections {
        bigint id PK
        string name
        string slug UK
        tinyint type
        string list_categories
        boolean is_active
    }
    settings {
        bigint id PK
        string key_col UK
        text value_col
        string type
    }
"""
    add_diagram(doc, 'Sơ đồ quan hệ thực thể (ERD)', erd, 'd15.png')
    
    doc.add_paragraph('Mô tả các mối quan hệ chính trong ERD:')
    add_table(doc, ['Quan hệ', 'Bảng 1', 'Bảng 2', 'Loại', 'Khóa ngoại', 'Mô tả'], [
        ('User - Order', 'users', 'orders', '1:N', 'orders.user_id → users.id', 'Một user có thể đặt nhiều đơn hàng. Xóa user → cascade xóa orders.'),
        ('User - Address', 'users', 'addresses', '1:N', 'addresses.user_id → users.id', 'Một user có nhiều địa chỉ giao hàng. Xóa user → cascade xóa addresses.'),
        ('Category - Product', 'categories', 'products', '1:N', 'products.category_id → categories.id', 'Một danh mục chứa nhiều sản phẩm. Xóa category → cascade xóa products.'),
        ('Category - Category', 'categories', 'categories', '1:N (self)', 'categories.parent_id → categories.id', 'Danh mục cha-con (self-referencing). Xóa cha → set null parent_id con.'),
        ('Product - ProductImage', 'products', 'product_images', '1:N', 'product_images.product_id → products.id', 'Một sản phẩm có nhiều ảnh gallery phụ. Xóa product → cascade xóa images.'),
        ('Product - OrderItem', 'products', 'order_items', '1:N', 'order_items.product_id → products.id', 'Một sản phẩm xuất hiện trong nhiều order items. Xóa product → cascade.'),
        ('Order - OrderItem', 'orders', 'order_items', '1:N', 'order_items.order_id → orders.id', 'Một đơn hàng chứa nhiều chi tiết sản phẩm. Xóa order → cascade xóa items.'),
        ('Order - Address', 'orders', 'addresses', 'N:1', 'orders.shipping_address_id → addresses.id', 'Nhiều đơn hàng có thể dùng chung một địa chỉ giao hàng.'),
    ])
    
    # --- Data Dictionary ---
    doc.add_heading('3.4.2. Từ điển dữ liệu (Data Dictionary)', level=3)
    doc.add_paragraph('Chi tiết cấu trúc từng bảng trong CSDL:')
    
    tables_dd = [
        ('Bảng users - Lưu trữ thông tin người dùng', [
            ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'Khóa chính'),
            ('first_name', 'VARCHAR(255)', 'NULLABLE', 'Tên'),
            ('last_name', 'VARCHAR(255)', 'NULLABLE', 'Họ'),
            ('email', 'VARCHAR(255)', 'UNIQUE, NOT NULL', 'Email đăng nhập'),
            ('username', 'VARCHAR(255)', 'UNIQUE, NULLABLE', 'Tên đăng nhập'),
            ('phone', 'VARCHAR(20)', 'NULLABLE', 'Số điện thoại'),
            ('password', 'VARCHAR(255)', 'NOT NULL', 'Mật khẩu (bcrypt)'),
            ('is_admin', 'BOOLEAN', 'DEFAULT false', 'true = Admin'),
            ('remember_token', 'VARCHAR(100)', 'NULLABLE', 'Token nhớ đăng nhập'),
        ]),
        ('Bảng categories - Danh mục sản phẩm (cha-con)', [
            ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'Khóa chính'),
            ('name', 'VARCHAR(255)', 'NOT NULL', 'Tên danh mục'),
            ('slug', 'VARCHAR(255)', 'UNIQUE', 'URL SEO-friendly'),
            ('description', 'TEXT', 'NULLABLE', 'Mô tả'),
            ('image', 'VARCHAR(255)', 'NULLABLE', 'Ảnh danh mục'),
            ('parent_id', 'BIGINT UNSIGNED', 'FK → categories.id, NULLABLE', 'ID danh mục cha'),
            ('is_active', 'BOOLEAN', 'DEFAULT true', 'Trạng thái'),
            ('order', 'INTEGER', 'DEFAULT 0', 'Thứ tự hiển thị'),
        ]),
        ('Bảng products - Sản phẩm', [
            ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'Khóa chính'),
            ('name', 'VARCHAR(255)', 'NOT NULL', 'Tên sản phẩm'),
            ('slug', 'VARCHAR(255)', 'UNIQUE', 'URL SEO-friendly'),
            ('description', 'TEXT', 'NULLABLE', 'Mô tả chi tiết'),
            ('price', 'DECIMAL(10,2)', 'NOT NULL', 'Giá bán (VNĐ)'),
            ('stock', 'INTEGER', 'DEFAULT 0', 'Tồn kho'),
            ('image', 'VARCHAR(255)', 'NULLABLE', 'Ảnh chính'),
            ('category_id', 'BIGINT UNSIGNED', 'FK → categories.id CASCADE', 'Danh mục'),
            ('is_active', 'BOOLEAN', 'DEFAULT true', 'Hiển thị'),
            ('is_featured', 'BOOLEAN', 'DEFAULT false', 'Nổi bật'),
        ]),
        ('Bảng orders - Đơn hàng', [
            ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'Khóa chính'),
            ('user_id', 'BIGINT UNSIGNED', 'FK → users.id CASCADE', 'Người đặt'),
            ('shipping_address_id', 'BIGINT UNSIGNED', 'FK → addresses.id', 'Địa chỉ giao'),
            ('order_number', 'VARCHAR(255)', 'UNIQUE', 'Mã đơn (OD...)'),
            ('status', 'ENUM', 'DEFAULT pending', 'pending/processing/shipping/completed/cancelled'),
            ('payment_status', 'ENUM', 'DEFAULT pending', 'pending/paid/failed/refunded'),
            ('payment_method', 'VARCHAR(255)', 'NULLABLE', 'cod/momo/bank_transfer'),
            ('payment_transaction_id', 'VARCHAR(255)', 'NULLABLE', 'Mã GD thanh toán'),
            ('payment_payload', 'JSON', 'NULLABLE', 'Dữ liệu từ cổng TT'),
            ('total_amount', 'DECIMAL(12,2)', 'NOT NULL', 'Tổng tiền'),
            ('notes', 'TEXT', 'NULLABLE', 'Ghi chú'),
            ('paid_at', 'TIMESTAMP', 'NULLABLE', 'Thời điểm thanh toán'),
        ]),
        ('Bảng order_items - Chi tiết đơn hàng', [
            ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'Khóa chính'),
            ('order_id', 'BIGINT UNSIGNED', 'FK → orders.id CASCADE', 'Đơn hàng'),
            ('product_id', 'BIGINT UNSIGNED', 'FK → products.id CASCADE', 'Sản phẩm'),
            ('quantity', 'INTEGER', 'NOT NULL', 'Số lượng mua'),
            ('price', 'DECIMAL(12,2)', 'NOT NULL', 'Đơn giá lúc mua'),
            ('subtotal', 'DECIMAL(12,2)', 'NOT NULL', 'Thành tiền'),
        ]),
        ('Bảng addresses - Địa chỉ giao hàng', [
            ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'Khóa chính'),
            ('user_id', 'BIGINT UNSIGNED', 'FK → users.id CASCADE', 'Người dùng'),
            ('full_name', 'VARCHAR(255)', 'NOT NULL', 'Họ tên người nhận'),
            ('phone', 'VARCHAR(255)', 'NOT NULL', 'SĐT người nhận'),
            ('address', 'VARCHAR(255)', 'NOT NULL', 'Địa chỉ chi tiết'),
            ('province_name', 'VARCHAR(255)', 'NULLABLE', 'Tỉnh/Thành phố'),
            ('district_name', 'VARCHAR(255)', 'NULLABLE', 'Quận/Huyện'),
            ('ward_name', 'VARCHAR(255)', 'NULLABLE', 'Phường/Xã'),
            ('is_default', 'BOOLEAN', 'DEFAULT false', 'Địa chỉ mặc định'),
        ]),
        ('Bảng product_images - Gallery ảnh SP', [
            ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'Khóa chính'),
            ('product_id', 'BIGINT UNSIGNED', 'FK → products.id CASCADE', 'Sản phẩm'),
            ('image_path', 'VARCHAR(255)', 'NOT NULL', 'Đường dẫn file ảnh'),
        ]),
        ('Bảng banners - Banner quảng cáo', [
            ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'Khóa chính'),
            ('title', 'VARCHAR(255)', 'NOT NULL', 'Tiêu đề'),
            ('image', 'VARCHAR(255)', 'NOT NULL', 'Ảnh banner'),
            ('link', 'VARCHAR(255)', 'NULLABLE', 'Liên kết'),
            ('order', 'INTEGER', 'DEFAULT 0', 'Thứ tự'),
            ('is_active', 'BOOLEAN', 'DEFAULT true', 'Trạng thái'),
        ]),
        ('Bảng home_sections - Section trang chủ', [
            ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'Khóa chính'),
            ('name', 'VARCHAR(255)', 'NOT NULL', 'Tên section'),
            ('slug', 'VARCHAR(255)', 'UNIQUE', 'URL-friendly'),
            ('type', 'TINYINT', 'NOT NULL', '1=Nổi bật, 2=Mới, 3=Theo DM'),
            ('list_categories', 'VARCHAR(255)', 'NULLABLE', 'ID danh mục (CSV)'),
            ('num', 'INTEGER', 'DEFAULT 8', 'Số SP hiển thị'),
            ('is_active', 'BOOLEAN', 'DEFAULT true', 'Trạng thái'),
        ]),
        ('Bảng settings - Cấu hình hệ thống', [
            ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'Khóa chính'),
            ('key', 'VARCHAR(255)', 'UNIQUE', 'Tên cấu hình'),
            ('value', 'TEXT', 'NULLABLE', 'Giá trị'),
            ('type', 'VARCHAR(255)', 'DEFAULT string', 'Kiểu dữ liệu'),
            ('description', 'VARCHAR(255)', 'NULLABLE', 'Mô tả'),
        ]),
    ]
    
    for title, cols in tables_dd:
        doc.add_paragraph('')
        doc.add_paragraph(title + ':', style='Intense Quote')
        add_table(doc, ['Tên cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Mô tả'], cols)
    
    # --- Diagram 16: Class Diagram ---
    doc.add_heading('3.4.3. Sơ đồ lớp (Class Diagram) - Quan hệ Model', level=3)
    doc.add_paragraph('Sơ đồ lớp thể hiện 7 Eloquent Model chính và các mối quan hệ giữa chúng (belongsTo, hasMany):')
    
    class_d = """
classDiagram
    class User {
        +bigint id
        +string email
        +boolean is_admin
        +addresses() HasMany
        +orders() HasMany
        +defaultAddress() Address
    }
    class Category {
        +bigint id
        +string name
        +string slug
        +parent() BelongsTo
        +children() HasMany
        +products() HasMany
    }
    class Product {
        +bigint id
        +string name
        +decimal price
        +category() BelongsTo
        +images() HasMany
        +orderItems() HasMany
    }
    class Order {
        +bigint id
        +string order_number
        +enum status
        +user() BelongsTo
        +items() HasMany
        +shippingAddress() BelongsTo
    }
    class OrderItem {
        +integer quantity
        +decimal price
        +order() BelongsTo
        +product() BelongsTo
    }
    class Address {
        +string full_name
        +boolean is_default
        +user() BelongsTo
    }
    class ProductImage {
        +string image_path
        +product() BelongsTo
    }
    User "1" --> "*" Order
    User "1" --> "*" Address
    Category "1" --> "*" Product
    Category "1" --> "*" Category
    Product "1" --> "*" OrderItem
    Product "1" --> "*" ProductImage
    Order "1" --> "*" OrderItem
    Order "*" --> "1" Address
"""
    add_diagram(doc, 'Sơ đồ lớp - Quan hệ Model Eloquent', class_d, 'd16.png')
    
    doc.add_paragraph('Mô tả các quan hệ trong Class Diagram:')
    add_table(doc, ['Model', 'Relationship', 'Loại', 'Mô tả'], [
        ('User', 'orders()', 'hasMany(Order)', 'Một user có nhiều đơn hàng.'),
        ('User', 'addresses()', 'hasMany(Address)', 'Một user có nhiều địa chỉ.'),
        ('User', 'defaultAddress()', 'Method', 'Trả về Address có is_default=true.'),
        ('Category', 'parent()', 'belongsTo(Category)', 'Danh mục con thuộc về danh mục cha (self-referencing).'),
        ('Category', 'children()', 'hasMany(Category)', 'Danh mục cha có nhiều danh mục con.'),
        ('Category', 'products()', 'hasMany(Product)', 'Một danh mục chứa nhiều sản phẩm.'),
        ('Product', 'category()', 'belongsTo(Category)', 'Một sản phẩm thuộc một danh mục.'),
        ('Product', 'images()', 'hasMany(ProductImage)', 'Một sản phẩm có nhiều ảnh gallery phụ.'),
        ('Product', 'orderItems()', 'hasMany(OrderItem)', 'Một sản phẩm xuất hiện trong nhiều chi tiết đơn hàng.'),
        ('Order', 'user()', 'belongsTo(User)', 'Một đơn hàng thuộc về một user.'),
        ('Order', 'items()', 'hasMany(OrderItem)', 'Một đơn hàng chứa nhiều chi tiết sản phẩm.'),
        ('Order', 'shippingAddress()', 'belongsTo(Address)', 'Một đơn hàng giao đến một địa chỉ.'),
        ('OrderItem', 'order()', 'belongsTo(Order)', 'Một chi tiết thuộc về một đơn hàng.'),
        ('OrderItem', 'product()', 'belongsTo(Product)', 'Một chi tiết liên kết đến một sản phẩm.'),
    ])
    
    # --- 3.5 Công nghệ ---
    doc.add_heading('3.5. Cơ sở lý thuyết về Công nghệ cốt lõi', level=2)
    doc.add_paragraph('Dự án sử dụng mô hình MVC (Model-View-Controller) trên nền tảng Laravel Framework:')
    add_table(doc, ['Công nghệ', 'Phiên bản', 'Vai trò'], [
        ('Laravel', '10.x/11.x', 'Framework PHP: Routing, Eloquent ORM, Blade Template, Auth, Middleware, Mail, Session, Validation.'),
        ('PHP', '8.1+', 'Ngôn ngữ backend.'),
        ('MySQL', '8.0+', 'CSDL quan hệ chính.'),
        ('Bootstrap', '5.3', 'CSS framework responsive.'),
        ('jQuery', '3.6', 'JavaScript: DOM, AJAX.'),
        ('Chart.js', 'Latest', 'Vẽ biểu đồ Dashboard.'),
        ('barryvdh/dompdf', 'Latest', 'Render HTML → PDF.'),
        ('Google Gemini API', 'Latest', 'AI chatbot tư vấn.'),
        ('MoMo Payment API', 'v2', 'Thanh toán ví MoMo.'),
        ('VietQR/NAPAS247', 'Latest', 'QR chuyển khoản ngân hàng.'),
        ('GHN API', 'v2', 'API địa chỉ Tỉnh/Huyện/Xã.'),
        ('SweetAlert2', '11.x', 'Popup/toast thông báo.'),
        ('Font Awesome', '6.x', 'Thư viện icon.'),
    ])
    
    # --- 3.6 Cấu trúc thư mục ---
    doc.add_heading('3.6. Cấu trúc thư mục dự án', level=2)
    add_table(doc, ['Thư mục / File', 'Mô tả'], [
        ('app/Http/Controllers/Admin/', '8 Controller quản trị: Dashboard, Category, Product, Order, User, Banner, HomeSection, Setting.'),
        ('app/Http/Controllers/Client/', '6 Controller khách hàng: Home, Product, Cart, Order, Payment, AiAssistant, User.'),
        ('app/Http/Controllers/Auth/', '2 Controller xác thực: Login (email/username), Register.'),
        ('app/Http/Controllers/', '2 Controller dùng chung: Address (CRUD + GHN proxy), Language (vi/en).'),
        ('app/Models/', '10 Eloquent Model: User, Category, Product, ProductImage, Order, OrderItem, Address, Banner, HomeSection, Setting.'),
        ('app/Services/', 'MomoService (thanh toán MoMo), GHNApiService (proxy API địa chỉ GHN).'),
        ('app/Mail/', 'OrderPlaced: email xác nhận đơn hàng.'),
        ('database/migrations/', '19 file migration định nghĩa schema CSDL.'),
        ('resources/views/admin/', 'Giao diện quản trị: dashboard, categories, products, orders, users, banners, home-sections, settings.'),
        ('resources/views/client/', 'Giao diện khách: home, products, cart, orders, addresses, payment, wishlist, user.'),
        ('resources/views/layouts/', 'Layout: eshopper.blade.php (client), admin.blade.php (admin).'),
        ('resources/lang/', 'Đa ngôn ngữ: vi/messages.php, en/messages.php.'),
        ('routes/web.php', 'Định tuyến URL: Public, Guest, Auth, Admin.'),
        ('public/js/wishlist.js', 'JavaScript xử lý Wishlist LocalStorage.'),
        ('public/css/theme.css', 'CSS tuỳ chỉnh giao diện.'),
    ])
    
    # --- 3.7 Module khó ---
    doc.add_heading('3.7. Giải pháp triển khai một số Module khó', level=2)
    
    doc.add_heading('3.7.1. Module Giỏ hàng bằng Session', level=3)
    doc.add_paragraph('Giỏ hàng được lưu trong Laravel Session dưới dạng mảng PHP associative array:')
    add_steps(doc, [
        'Cấu trúc dữ liệu: session("cart") = [product_id => ["name" => "...", "price" => 50000, "quantity" => 2, "image" => "..."], ...].',
        'Thêm SP (CartController@add): Kiểm tra product_id đã tồn tại → nếu có: cộng quantity, nếu chưa: thêm key mới.',
        'Cập nhật (CartController@update): Thay đổi quantity của product_id chỉ định trong session.',
        'Xóa SP (CartController@remove): Dùng unset() xóa key product_id khỏi mảng cart.',
        'Checkout (CartController@checkout): Load mảng cart + top 3 địa chỉ user → render form thanh toán.',
        'Khi đặt hàng thành công: session()->forget("cart") để xóa toàn bộ giỏ hàng.',
        'Ưu điểm: không cần bảng DB riêng cho cart, đơn giản, nhanh. Nhược điểm: mất khi hết session.',
    ])
    
    doc.add_heading('3.7.2. Module Wishlist bằng LocalStorage', level=3)
    doc.add_paragraph('Sản phẩm yêu thích lưu hoàn toàn ở phía client (trình duyệt), xử lý bằng file wishlist.js:')
    add_steps(doc, [
        'LocalStorage key: "venshop_wishlist", value: JSON array chứa các object {id, name, price, image, url}.',
        'toggleWishlist(id, name, price, image, url): Kiểm tra SP đã lưu chưa → nếu có: splice() xóa khỏi mảng, nếu chưa: push() thêm vào.',
        'updateBadge(): Cập nhật số lượng trên badge icon trái tim ở header. Nếu wishlist.length > 0 → hiện badge, ngược lại → ẩn.',
        'updateHeartIcons(): Quét tất cả nút .btn-wishlist trên trang, so sánh data-id với wishlist → đổi icon "far fa-heart" (rỗng) thành "fas fa-heart text-danger" (đỏ đậm) hoặc ngược lại.',
        'renderWishlistPage(): Render danh sách SP yêu thích dạng grid card trên trang /wishlist. Hiển thị nút X để xóa từng SP.',
        'showToast(): Sử dụng SweetAlert2 Toast hiển thị thông báo "Đã thêm/bỏ SP yêu thích".',
        'Ưu điểm: không cần đăng nhập, không tạo bảng DB, giảm tải server. Nhược điểm: mất khi xóa dữ liệu trình duyệt.',
    ])
    
    doc.add_heading('3.7.3. Tích hợp thanh toán MoMo', level=3)
    doc.add_paragraph('Quy trình kỹ thuật tích hợp thanh toán ví MoMo:')
    add_steps(doc, [
        'MomoService xây dựng request body: partnerCode, accessKey, amount, orderId, requestId, orderInfo, returnUrl (/payment/momo/return), ipnUrl (/payment/momo/ipn).',
        'Tạo chữ ký HMAC SHA256: rawSignature = "accessKey=...&amount=...&..." → hash_hmac("sha256", rawSignature, secretKey).',
        'POST request đến MoMo API endpoint → nhận response chứa payUrl (link thanh toán) và qrCodeUrl (ảnh QR code).',
        'Lưu toàn bộ response vào cột payment_payload (JSON) của order để truy vết sau này.',
        'Xử lý IPN callback (momoIpn): MoMo gọi POST server-to-server → verify lại signature → nếu resultCode=0: payment_status="paid" + paid_at=now().',
        'Xử lý returnUrl (momoReturn): Redirect user trở lại website → hiển thị kết quả thanh toán.',
        'momoMockSuccess (chỉ dev): Endpoint giả lập thanh toán thành công để test.',
    ])
    
    doc.add_heading('3.7.4. Tích hợp chuyển khoản VietQR', level=3)
    doc.add_paragraph('Quy trình tạo mã QR thanh toán chuyển khoản tự động:')
    add_steps(doc, [
        'Đọc thông tin ngân hàng từ bảng settings: bank_bin (mã BIN), bank_account_no, bank_account_name.',
        'Tạo URL QR VietQR: https://img.vietqr.io/image/{bank_bin}-{account_no}-print.png?amount={total_amount}&addInfo={order_number}.',
        'QR code chứa: số tài khoản, số tiền, nội dung chuyển khoản = mã đơn hàng → khách quét QR bằng app NH → thông tin tự điền.',
        'Khách click "Tôi đã chuyển khoản" → bankNotify(): ghi timestamp xác nhận vào payment_payload.',
        'Admin kiểm tra tài khoản NH thực tế → click "Xác nhận thanh toán" (markPaid) → ghi admin_id + transaction_id vào payment_payload, payment_status="paid".',
    ])
    
    doc.add_heading('3.7.5. Trợ lý AI tư vấn (AiAssistantController)', level=3)
    doc.add_paragraph('Kiến trúc chatbot AI tư vấn sản phẩm sử dụng Google Gemini:')
    add_steps(doc, [
        'extractKeywords(prompt): Tách từ khóa từ câu hỏi, loại bỏ stop words tiếng Việt (tôi, muốn, cần, có, ...) → mảng keywords.',
        'findRelatedProducts(keywords): Query bảng products với WHERE name LIKE %keyword1% OR name LIKE %keyword2%..., chỉ lấy SP active, limit 10.',
        'getFeaturedProducts(): Lấy thêm 5 SP nổi bật (is_featured=true) để bổ sung gợi ý.',
        'getCategorySummary(): Query bảng categories lấy tên + số lượng SP mỗi danh mục → cung cấp ngữ cảnh cho AI.',
        'buildSystemInstruction(): Xây dựng system prompt chứa: tên cửa hàng, vai trò (tư vấn viên), danh sách danh mục, danh sách SP thực (tên + giá + URL).',
        'callGeminiWithFallback(): Gọi Google Gemini API với cơ chế fallback: thử model chính (gemini-2.0-flash) → nếu lỗi → thử model dự phòng (gemini-1.5-flash).',
        'buildContents(): Xây dựng mảng messages gồm system instruction + lịch sử chat (history) + câu hỏi mới → gửi API.',
        'Response: Trả JSON chứa reply (text AI trả lời) + suggested_products (mảng SP gợi ý với name, price, url).',
        'Frontend: Floating panel draggable (kéo thả được), render markdown, hiển thị card SP gợi ý, lịch sử lưu LocalStorage (20 lượt/phiên).',
    ])
    
    doc.add_heading('3.7.6. Bảo mật thông tin người dùng trong Admin', level=3)
    doc.add_paragraph('Giải pháp che giấu thông tin nhạy cảm trên trang quản lý người dùng:')
    add_steps(doc, [
        'Server-side masking (Blade PHP): Email: preg_replace("/(?<=.{2}).(?=[^@]*?@)/", "*", $email) → hiện 2 ký tự đầu, che phần còn lại trước @. SĐT: substr(0,3) + str_repeat("*", len-6) + substr(-3) → hiện 3 đầu + 3 cuối.',
        'HTML structure: Mỗi ô chứa 2 span: .masked-text (hiển thị bản che) và .full-text.d-none (ẩn bản đầy đủ). Kèm icon <i class="fas fa-eye toggle-eye">.',
        'JavaScript toggle: Click icon mắt → toggle class d-none giữa masked-text và full-text, đổi icon fa-eye ↔ fa-eye-slash.',
        'Kết quả: Admin thấy "cu*****@gmail.com" và "090****123" mặc định. Bấm icon mắt → hiện đầy đủ. Bấm lại → ẩn.',
    ])
    
    # --- Diagram 17: Deployment ---
    doc.add_heading('3.8. Sơ đồ triển khai (Deployment Diagram)', level=2)
    doc.add_paragraph('Sơ đồ triển khai thể hiện cách các thành phần hệ thống được phân bổ trên các node vật lý/logic:')
    
    deploy = """
graph LR
    subgraph User [Thiet bi Nguoi dung]
        B1[Desktop Browser]
        B2[Mobile Browser]
    end
    subgraph Server [Web Server]
        WEB[Apache / Nginx]
        APP[Laravel Application]
        STORE[File Storage]
    end
    subgraph DB [Database Server]
        MYSQL[(MySQL 8.0)]
    end
    subgraph Ext [Dich vu ben ngoai]
        MOMO[MoMo Payment]
        GHN[GHN Logistics]
        GEMINI[Google Gemini AI]
        VIETQR[VietQR NAPAS247]
        SMTP[SMTP Mail]
    end
    B1 -- HTTPS --> WEB
    B2 -- HTTPS --> WEB
    WEB --> APP
    APP --> MYSQL
    APP --> STORE
    APP --> MOMO
    APP --> GHN
    APP --> GEMINI
    APP --> VIETQR
    APP --> SMTP
"""
    add_diagram(doc, 'Sơ đồ triển khai hệ thống', deploy, 'd17.png')
    
    doc.add_paragraph('Mô tả chi tiết từng thành phần trong sơ đồ triển khai:')
    add_table(doc, ['Thành phần', 'Vai trò', 'Công nghệ / Giao thức'], [
        ('Desktop/Mobile Browser', 'Giao diện người dùng cuối, hiển thị trang web responsive.', 'HTML/CSS/JS, Bootstrap 5, jQuery, LocalStorage.'),
        ('Apache/Nginx', 'Web server xử lý HTTP request, phục vụ file tĩnh, proxy pass đến PHP-FPM.', 'HTTP/HTTPS, SSL/TLS.'),
        ('Laravel Application', 'Application server xử lý business logic theo mô hình MVC.', 'PHP 8.1+, Laravel 10.x/11.x, Composer.'),
        ('File Storage', 'Lưu trữ file upload (ảnh sản phẩm, danh mục, banner, logo).', 'public/storage (symlink đến storage/app/public).'),
        ('MySQL 8.0', 'Lưu trữ toàn bộ dữ liệu nghiệp vụ (10 bảng, 19 migration).', 'SQL, PDO driver, InnoDB engine.'),
        ('MoMo Payment', 'Cổng thanh toán ví điện tử MoMo, xử lý QR payment + IPN webhook.', 'REST API v2, HMAC SHA256 signature.'),
        ('GHN Logistics', 'API lấy danh sách Tỉnh/Quận/Phường cho form địa chỉ.', 'REST API, Token authentication.'),
        ('Google Gemini AI', 'API trí tuệ nhân tạo cho chatbot tư vấn sản phẩm.', 'REST API, API Key, model: gemini-2.0-flash.'),
        ('VietQR/NAPAS247', 'Tạo mã QR thanh toán chuyển khoản ngân hàng động.', 'Image URL API, chuẩn NAPAS247.'),
        ('SMTP Mail', 'Gửi email xác nhận đơn hàng tự động.', 'SMTP protocol, Laravel Mail + Mailable class.'),
    ])
    
    # Save
    output_path = '/Users/nguyennghia/Desktop/workspace/eco-web-laravel/TaiLieu_Chuong2_Chuong3.docx'
    doc.save(output_path)
    print(f"\n=== Document saved to: {output_path} ===")
    print(f"Total diagrams: 17")

if __name__ == '__main__':
    main()
